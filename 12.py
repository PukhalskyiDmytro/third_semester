import os
from datetime import date
from typing import List
from docx import Document
from openpyxl import load_workbook

def transliterate(name: str) -> str:
    mapping = {
        'А':'A','Б':'B','В':'V','Г':'H','Ґ':'G','Д':'D','Е':'E','Є':'Ye','Ж':'Zh','З':'Z','И':'Y','І':'I','Ї':'Yi','Й':'Y','К':'K','Л':'L','М':'M','Н':'N','О':'O','П':'P','Р':'R','С':'S','Т':'T','У':'U','Ф':'F','Х':'Kh','Ц':'Ts','Ч':'Ch','Ш':'Sh','Щ':'Shch','Ь':'','Ю':'Yu','Я':'Ya',
        'а':'a','б':'b','в':'v','г':'h','ґ':'g','д':'d','е':'e','є':'ie','ж':'zh','з':'z','и':'y','і':'i','ї':'i','й':'i','к':'k','л':'l','м':'m','н':'n','о':'o','п':'p','р':'r','с':'s','т':'t','у':'u','ф':'f','х':'kh','ц':'ts','ч':'ch','ш':'sh','щ':'shch','ь':'','ю':'iu','я':'ia'
    }
    out = []
    for ch in name:
        if ch in mapping:
            out.append(mapping[ch])
        else:
            if ord(ch) < 128:
                out.append(ch)
            else:
                pass
    s = ''.join(out)
    s = s.replace(' ', '_')
    s = ''.join(c for c in s if c.isalnum() or c == '_')
    return s

def replace_in_doc(doc: Document, placeholder: str, value: str):
    for p in doc.paragraphs:
        for r in p.runs:
            if placeholder in r.text:
                r.text = r.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if placeholder in r.text:
                            r.text = r.text.replace(placeholder, value)

def generate_reports(template_path: str = 'lab_report_template.docx', excel_path: str = 'lab_results.xlsx'):
    wb = load_workbook(excel_path, data_only=True)
    if 'Labs' in wb.sheetnames:
        ws = wb['Labs']
    else:
        ws = wb[wb.sheetnames[0]]
    headers = {}
    for col in range(1, ws.max_column+1):
        h = ws.cell(row=1, column=col).value
        if h is None:
            continue
        headers[str(h).strip()] = col
    lab_cols = [headers[k] for k in headers if k.startswith('Lab')]
    lab_cols.sort()
    labs_total = len(lab_cols)
    avg_col = headers.get('Average')
    out_dir = 'lab_reports'
    os.makedirs(out_dir, exist_ok=True)
    for row in range(2, ws.max_row+1):
        full = ws.cell(row=row, column=headers.get('FullName')).value if headers.get('FullName') else ''
        group = ws.cell(row=row, column=headers.get('Group')).value if headers.get('Group') else ''
        if not full:
            continue
        scores: List[float] = []
        for c in lab_cols:
            v = ws.cell(row=row, column=c).value
            if v is None or (isinstance(v, str) and str(v).strip() == ''):
                continue
            try:
                num = float(v)
            except Exception:
                continue
            scores.append(num)
        average = None
        if avg_col:
            a = ws.cell(row=row, column=avg_col).value
            try:
                if a is not None and str(a).strip() != '':
                    average = float(a)
            except Exception:
                average = None
        if average is None:
            if scores:
                average = sum(scores)/len(scores)
            else:
                average = 0.0
        labs_done = len(scores)
        result_text = 'виконав(ла) всі лабораторні роботи на належному рівні' if average >= 70 and labs_done == labs_total and labs_total>0 else 'потребує додаткового опрацювання матеріалу'
        doc = Document(template_path)
        replace_in_doc(doc, '{{FULL_NAME}}', str(full))
        replace_in_doc(doc, '{{GROUP}}', str(group))
        replace_in_doc(doc, '{{LABS_DONE}}', str(labs_done))
        replace_in_doc(doc, '{{LABS_TOTAL}}', str(labs_total))
        replace_in_doc(doc, '{{AVERAGE_SCORE}}', f"{average:.2f}")
        replace_in_doc(doc, '{{RESULT_TEXT}}', result_text)
        replace_in_doc(doc, '{{DATE}}', date.today().isoformat())
        fname = f"lab_report_{transliterate(str(full))}.docx"
        path = os.path.join(out_dir, fname)
        doc.save(path)

if __name__ == '__main__':
    generate_reports()
